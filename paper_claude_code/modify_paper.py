#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修改论文，添加图表和AI技术内容
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 创建新文档
doc = Document()

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# 设置字体
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('本科毕业论文')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 论文标题
title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run('城市商业银行信用风险管理研究')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

title3 = doc.add_paragraph()
title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title3.add_run('------以邯郸银行为例')
run.bold = True
run.font.size = Pt(16)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

# 学校信息
info_items = [
    ('学 院：', '金融学院'),
    ('专 业：', '金融学'),
    ('学生姓名：', ''),
    ('学 号：', ''),
    ('指导教师：', ''),
    ('完成时间：', '2025年'),
]

for prefix, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(prefix + value)
    run.font.size = Pt(14)

doc.add_paragraph()

# 摘要标题
abstract_title = doc.add_paragraph()
abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = abstract_title.add_run('摘 要')
run.bold = True
run.font.size = Pt(14)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 摘要内容
abstract_content = doc.add_paragraph()
abstract_text = '''城市商业银行是我国银行体系的重要组成部分，承担着服务地方经济、支持中小微企业融资的重要职能。然而，由于区域集中度高、客户资质参差不齐、风险分散能力相对有限，信用风险始终是城商行面临的核心挑战。党的二十大以来，金融高质量发展被提升至战略高度，进一步强化了城商行完善信用风险管理体系的紧迫性。

本文以邯郸银行股份有限公司为研究对象，系统梳理该行2020---2024年五年间的年度报告数据，运用信息不对称理论、全面风险管理理论及巴塞尔协议框架对其信用风险管理实践进行深入分析。研究发现：邯郸银行信用风险管理取得了一定成效，资产规模从2020年末1,809亿元增至2024年末2,492亿元，不良贷款率由2022年高点1.90%攀升至2023年2.24%后于2024年显著回落至1.38%，拨备覆盖率于2024年提升至204.35%，资本充足率持续高于监管要求；但同时也存在信贷结构高度集中于制造业与批发零售业、关注类贷款占比持续攀升（2023年达10.07%）、风险识别技术体系较为落后、风控执行机制存在松弛现象、信用风险人才相对匮乏等突出问题。

针对上述问题，本文提出优化信贷全流程管理、升级智能风控系统、强化执行层面管理与推进人才能力建设四个维度的对策建议，旨在为邯郸银行及同类城市商业银行提升信用风险管理水平提供参考。

关键词：城市商业银行；信用风险；风险管理；邯郸银行；不良贷款率'''
abstract_content.add_run(abstract_text)
abstract_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Abstract
doc.add_paragraph()
abstract_title_en = doc.add_paragraph()
abstract_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = abstract_title_en.add_run('Abstract')
run.bold = True
run.font.size = Pt(14)

abstract_content_en = doc.add_paragraph()
abstract_text_en = '''City commercial banks play a vital role in China's banking system by serving local economies and supporting small-to-medium enterprises. Due to their high regional concentration, heterogeneous borrower quality, and limited risk diversification capacity, credit risk remains their core challenge. This paper takes Handan Bank Co., Ltd. as a case study, systematically analyzes its annual report data from 2020 to 2024, and applies theoretical frameworks including information asymmetry, enterprise risk management, and the Basel Accord to evaluate its credit risk management practices. The research finds that while Handan Bank has achieved notable progress---total assets growing from CNY 180.9 billion in 2020 to CNY 249.2 billion in 2024, and non-performing loan (NPL) ratio declining sharply from 2.24% in 2023 to 1.38% in 2024---persistent problems remain, including a concentrated credit portfolio, rising special-mention loans (peaking at 10.07% of total loans in 2023), outdated risk identification technology, lapses in risk control execution, and insufficient risk management talent. The paper proposes recommendations across four dimensions: optimizing the full credit life cycle, upgrading intelligent risk control systems, strengthening execution management, and building talent capacity.

Keywords: city commercial bank; credit risk; risk management; Handan Bank; non-performing loan ratio'''
abstract_content_en.add_run(abstract_text_en)
abstract_content_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 第一章
doc.add_page_break()
ch1_title = doc.add_paragraph()
ch1_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ch1_title.add_run('一、引言')
run.bold = True
run.font.size = Pt(16)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# ... (继续添加原有论文内容，为了简化，这里继续添加内容)
# 由于原论文很长，我会继续添加关键部分

# 1.1 研究背景
section_title = doc.add_paragraph()
section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = section_title.add_run('（一）研究背景')
run.bold = True
run.font.size = Pt(14)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

para = doc.add_paragraph()
para.add_run('''金融是现代经济的核心，银行业的稳健运行对于维护经济金融秩序具有不可替代的基础性作用。作为我国银行体系的重要组成部分，城市商业银行（以下简称"城商行"）自20世纪90年代城市信用社改制以来，历经三十余年发展，资产规模持续扩张，已成为地方经济发展不可或缺的金融支柱。根据中国银行保险监督管理委员会披露数据，截至2024年末，全国城商行总资产规模已突破60万亿元，在银行业整体资产中占比约15%。

然而，与大型国有商业银行及全国性股份制银行相比，城商行在信用风险管理方面面临更为复杂的挑战。一方面，城商行客户群体以中小企业为主，这类客户财务透明度低、抗风险能力弱、信息不对称程度高；另一方面，城商行信贷资产高度集中于特定区域和行业，风险分散效果有限。近年来，受宏观经济下行压力加大、房地产行业调整、地方政府债务风险显现等多重因素叠加影响，部分城商行不良贷款余额和不良贷款率呈现双升态势，信用风险管理面临严峻考验。

邯郸银行股份有限公司（以下简称"邯郸银行"）是河北省邯郸市最大的地方法人金融机构，成立于2005年，前身为1998年组建的邯郸市商业银行。该行深耕邯郸区域，业务辐射石家庄、保定、邢台、秦皇岛等地，是典型的区域性城商行。从历年报告来看，邯郸银行资产规模快速扩张：2020年末总资产1,809亿元，2024年末已达2,492亿元，年均复合增长率约8.3%；同期贷款余额从760亿元增长至1,385亿元。但与此同时，不良贷款率在2022---2023年间出现明显攀升（从1.90%升至2.24%），关注类贷款占比更于2023年末高达10.07%，潜在信用风险值得高度重视。2024年度，邯郸银行通过加强不良资产处置，不良贷款率大幅回落至1.38%，拨备覆盖率提升至204.35%，风险管理能力有所增强，但深层次问题仍然存在。

在此背景下，深入分析邯郸银行信用风险管理的现状、问题与改进路径，不仅对该行自身具有重要的实践价值，对于同类城商行提升信用风险管理水平也具有重要的参考意义。''')
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 1.2 研究意义
section_title = doc.add_paragraph()
section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = section_title.add_run('（二）研究意义')
run.bold = True
run.font.size = Pt(14)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

para = doc.add_paragraph()
para.add_run('1. 理论意义\n\n')
run = para.add_run('')
run = para.add_run('本文将信息不对称理论、全面风险管理理论和巴塞尔协议框架应用于城商行信用风险管理分析，丰富了银行信用风险管理理论在区域性中小银行场景下的应用研究。通过对邯郸银行五年连续数据的纵向分析，本文尝试识别城商行信用风险管理的共性规律与特殊矛盾，对于完善城商行信用风险管理理论体系具有积极贡献。')

para = doc.add_paragraph()
para.add_run('2. 实践意义\n\n')
run = para.add_run('')
run = para.add_run('本文研究成果可为邯郸银行优化信贷结构、健全风险预警机制、强化内部控制提供具体的对策建议，有助于该行在业务扩张的同时保持资产质量稳定。同时，本文所总结的经验与教训对于河北省及全国同类城商行防范化解信用风险、支持实体经济发展具有借鉴价值，也可为金融监管部门制定差异化监管政策提供参考。')

# 由于完整论文很长，我需要继续添加内容
# 让我创建一个更完整的版本

print("论文基础结构已创建")
print("正在继续添加完整内容...")

# 保存当前进度并创建完整版本
doc.save('handan-bank-risk-paper-v2_temp.docx')
print("临时文件已保存")

# 现在我需要创建包含完整内容的论文
# 由于完整论文很长，我会用另一种方式处理
