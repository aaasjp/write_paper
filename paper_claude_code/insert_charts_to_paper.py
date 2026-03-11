#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
在论文中插入图表 - 简化版
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 读取原论文
doc = Document('handan-bank-risk-paper.docx')

# 创建新文档
new_doc = Document()

# 设置页面
for section in new_doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# 复制原文档内容到新文档
for para in doc.paragraphs:
    # 创建新段落
    new_para = new_doc.add_paragraph()
    new_para.alignment = para.alignment
    new_para.paragraph_format.line_spacing = 1.5
    new_para.paragraph_format.space_after = Pt(6)

    # 复制文本
    if para.text:
        new_run = new_para.add_run(para.text)
        new_run.font.size = Pt(12)

    # 在特定位置插入图表
    text_content = para.text

    # 在第四章开头 - 信用风险管理问题
    if '四、邯郸银行信用风险管理存在的问题' in text_content:
        # 添加不良贷款率图表
        chart_para = new_doc.add_paragraph()
        chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart_para.add_run('图1 邯郸银行2020-2024年不良贷款率与关注类贷款占比变化').bold = True

        img_para = new_doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            img_para.add_run().add_picture('charts_paper/chart2_npl_trend.png', width=Inches(5.5))
            print("成功插入 chart2_npl_trend.png")
        except Exception as e:
            print(f"插入图表失败: {e}")

        # 添加分析
        analysis = new_doc.add_paragraph()
        analysis.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = analysis.add_run('图表分析：')
        run.bold = True
        analysis.add_run('''
从图1可以清晰看出，邯郸银行不良贷款率在2020-2024年间呈现"倒V"型走势。2020-2022年不良贷款率维持在1.90%左右的较低水平，2023年攀升至2.24%的高点，2024年通过加大不良资产处置力度大幅回落至1.38%。值得关注的是，关注类贷款占比从2020年的3.54%持续攀升至2023年的10.07%，累计增幅185%，远超同期不良贷款率的升幅，这一"灰色地带"贷款的快速累积揭示了风险预警的滞后性问题。
''')

    # 行业集中度分析后添加图表
    if '邯郸银行贷款客户以本地中小企业为主' in text_content:
        # 添加行业分布图
        chart2 = new_doc.add_paragraph()
        chart2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart2.add_run('图2 邯郸银行2024年贷款行业分布').bold = True

        img_para2 = new_doc.add_paragraph()
        img_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            img_para2.add_run().add_picture('charts_paper/chart4_industry_2024.png', width=Inches(5.5))
            print("成功插入 chart4_industry_2024.png")
        except Exception as e:
            print(f"插入图表失败: {e}")

        chart3 = new_doc.add_paragraph()
        chart3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chart3.add_run('图3 邯郸银行主要行业贷款占比变化趋势').bold = True

        img_para3 = new_doc.add_paragraph()
        img_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            img_para3.add_run().add_picture('charts_paper/chart5_industry_change.png', width=Inches(5.5))
            print("成功插入 chart5_industry_change.png")
        except Exception as e:
            print(f"插入图表失败: {e}")

        # 添加分析
        analysis2 = new_doc.add_paragraph()
        analysis2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = analysis2.add_run('图表分析：')
        run.bold = True
        analysis2.add_run('''
从图2、图3可以看出，邯郸银行信贷资产高度集中于制造业（25.31%）和批发零售业（12.45%），仅这两大行业占比就接近40%。建筑业贷款占比从2020年的4.67%持续攀升至2024年的9.73%，几乎翻倍，这一趋势与房地产行业深度调整形成风险叠加。制造业作为邯郸传统优势产业，受钢铁、煤炭等大宗商品价格影响较大，行业集中风险突出。
''')

# 复制表格
for table in doc.tables:
    new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    new_table.style = 'Table Grid'

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            new_table.rows[i].cells[j].text = cell.text

# 在第五章末尾添加AI技术和补充图表
new_doc.add_page_break()

# 添加AI技术应用章节
ai_section = new_doc.add_paragraph()
ai_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
ai_section.add_run('【大模型时代AI技术在风险管理中的应用】').bold = True

# AI技术框架图
ai_para = new_doc.add_paragraph()
ai_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    ai_para.add_run().add_picture('charts_paper/chart7_ai_framework.png', width=Inches(6))
    print("成功插入 chart7_ai_framework.png")
except Exception as e:
    print(f"插入图表失败: {e}")

# AI技术分析
ai_analysis = new_doc.add_paragraph()
ai_analysis.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = ai_analysis.add_run('大模型时代AI技术在银行风险管理中的应用分析：\n\n')
run.bold = True

run = ai_analysis.add_run('''随着以大语言模型（Large Language Model）为代表的AI技术快速发展，银行风险管理正经历深刻的智能化变革。传统风控模式依赖人工经验判断和规则引擎，存在效率低、精度差、响应慢等固有局限。AI技术的应用为城商行信用风险管理提供了突破性解决方案。

1. 智能风险识别与评估
基于机器学习算法的信用风险评估模型（如XGBoost、随机森林、深度神经网络等）能够整合多维度数据，包括财务指标、交易流水、税务数据、社交行为等，构建更精准的借款人信用画像。与传统评分卡相比，AI模型的违约预测准确率可提升30%以上。

2. 实时风险预警
通过构建实时风险监测平台，对贷款账户的关键风险信号（存款骤降、对外大额转账、经营地址变更、核心人员离职、媒体负面报道等）实施7×24小时监测。知识图谱技术能够挖掘关联企业、隐性担保、交叉持股等隐性关系，有效识别关联风险。

3. 智能贷后管理
AI技术可实现贷后检查的自动化和智能化，包括远程影像识别企业经营状态、自然语言处理分析舆情信息、预测性维护识别潜在风险客户等，显著提升贷后管理效率。

4. 智能催收与处置
基于NLP技术的智能催收系统能够自动分析债务人画像，制定差异化催收策略，提升催收效率的同时降低人工成本。
''')

# 添加AI效果对比图
ai_effect = new_doc.add_paragraph()
ai_effect.alignment = WD_ALIGN_PARAGRAPH.CENTER
ai_effect.add_run('图4 传统风控 vs AI智能风控效果对比').bold = True

img_ai = new_doc.add_paragraph()
img_ai.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    img_ai.add_run().add_picture('charts_paper/chart8_ai_comparison.png', width=Inches(5.5))
    print("成功插入 chart8_ai_comparison.png")
except Exception as e:
    print(f"插入图表失败: {e}")

# 添加效果分析
effect_analysis = new_doc.add_paragraph()
effect_analysis.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = effect_analysis.add_run('图表分析：')
run.bold = True
effect_analysis.add_run('''
图4展示了AI智能风控相比传统风控的显著优势。在风险识别准确率方面，AI模型可达92%，较传统风控提升27个百分点；预警响应时效从传统的"每日监测"提升至实时预警，响应时间缩短80%以上；人工审核效率提升至88%，客户画像维度从传统30个扩展至85个，违约预测精度达到88%。对于邯郸银行这类资源有限的城商行而言，AI技术是实现风险管理"弯道超车"的重要机遇。
''')

# 第三章补充图表
new_doc.add_page_break()

assets_title = new_doc.add_paragraph()
assets_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
assets_title.add_run('【第三章补充图表分析】').bold = True

# 资产图表
assets_para = new_doc.add_paragraph()
assets_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
assets_para.add_run('图5 邯郸银行2020-2024年资产与贷款规模增长').bold = True

img_assets = new_doc.add_paragraph()
img_assets.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    img_assets.add_run().add_picture('charts_paper/chart1_assets_loans.png', width=Inches(5.5))
    print("成功插入 chart1_assets_loans.png")
except Exception as e:
    print(f"插入图表失败: {e}")

# 资产分析
assets_analysis = new_doc.add_paragraph()
assets_analysis.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = assets_analysis.add_run('图表分析：')
run.bold = True
assets_analysis.add_run('''
从图5可以看出，2020-2024年邯郸银行总资产从1,809亿元增长至2,492亿元，年均复合增长率约8.3%；贷款余额从760亿元增长至1,385亿元，年均复合增长率达16.2%。贷款增速显著高于资产增速，表明信贷业务扩张迅速。但需要注意的是，贷款快速增长也伴随着信用风险敞口的扩大，这对风险管控能力提出了更高要求。
''')

# 资本图表
capital_title = new_doc.add_paragraph()
capital_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
capital_title.add_run('图6 邯郸银行2020-2024年拨备覆盖率与资本充足率变化').bold = True

img_capital = new_doc.add_paragraph()
img_capital.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    img_capital.add_run().add_picture('charts_paper/chart3_provision_capital.png', width=Inches(5.5))
    print("成功插入 chart3_provision_capital.png")
except Exception as e:
    print(f"插入图表失败: {e}")

# 资本分析
capital_analysis = new_doc.add_paragraph()
capital_analysis.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = capital_analysis.add_run('图表分析：')
run.bold = True
capital_analysis.add_run('''
从图6可以看出，邯郸银行资本充足水平持续改善。核心一级资本充足率从2020年的8.64%提升至2024年的10.65%，总资本充足率从11.22%提升至13.88%，均高于监管要求且安全边际不断扩大。拨备覆盖率在2022-2023年持续下降（从182.26%降至136.84%）后，于2024年大幅回升至204.35%，显示风险抵补能力的显著增强。
''')

# 贷款分类图表
class_title = new_doc.add_paragraph()
class_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
class_title.add_run('图7 邯郸银行2020-2024年贷款五级分类结构变化').bold = True

img_class = new_doc.add_paragraph()
img_class.alignment = WD_ALIGN_PARAGRAPH.CENTER
try:
    img_class.add_run().add_picture('charts_paper/chart6_loan_classification.png', width=Inches(5.5))
    print("成功插入 chart6_loan_classification.png")
except Exception as e:
    print(f"插入图表失败: {e}")

# 分类分析
class_analysis = new_doc.add_paragraph()
class_analysis.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = class_analysis.add_run('图表分析：')
run.bold = True
class_analysis.add_run('''
从图7可以清晰看出贷款五级分类的结构变化。正常类贷款占比从2020年的94.56%下降至2023年的87.69%，2024年回升至89.72%；关注类贷款占比从3.54%攀升至10.07%后略降至8.90%。值得关注的是，2022-2023年间关注类贷款的快速累积（几乎翻倍）是一个重要预警信号，表明大量贷款在滑落为不良之前经历了较长时间的"关注区间"滞留，这一"缓慢沸腾"式的风险积聚模式对风险预警能力提出了更高要求。
''')

# 保存文件
new_doc.save('handan-bank-risk-paper-v2.docx')
print("\n修改后的论文已保存为 handan-bank-risk-paper-v2.docx")