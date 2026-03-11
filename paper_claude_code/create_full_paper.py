#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建完整的论文，包含图表和AI技术内容
使用python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement(f'w:shd [{nsdecls("w")}]')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_row(table, cells_text, bold=False, font_size=10, bg_color=None):
    """添加表格行"""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run()
        run.font.size = Pt(font_size)
        run.bold = bold
        if bg_color and i == 0:
            set_cell_shading(cell, bg_color)
    return row

# 创建文档
doc = Document()

# 设置页面 (A4)
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# 设置默认字体样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ========== 标题页 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('本 科 毕 业 论 文')
run.bold = True
run.font.size = Pt(22)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run('城市商业银行信用风险管理研究')
run.bold = True
run.font.size = Pt(20)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

title3 = doc.add_paragraph()
title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title3.add_run('——以邯郸银行为例')
run.bold = True
run.font.size = Pt(16)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
doc.add_paragraph()

# 学校信息
for info in ['学    院：金融学院', '专    业：金融学', '学生姓名：', '学    号：', '指导教师：', '完成时间：2025年']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(info)
    run.font.size = Pt(14)

doc.add_page_break()

# ========== 摘要 ==========
abstract_title = doc.add_paragraph()
abstract_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = abstract_title.add_run('摘  要')
run.bold = True
run.font.size = Pt(14)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

abstract_text = '''城市商业银行是我国银行体系的重要组成部分，承担着服务地方经济、支持中小微企业融资的重要职能。然而，由于区域集中度高、客户资质参差不齐、风险分散能力相对有限，信用风险始终是城商行面临的核心挑战。党的二十大以来，金融高质量发展被提升至战略高度，进一步强化了城商行完善信用风险管理体系的紧迫性。

本文以邯郸银行股份有限公司为研究对象，系统梳理该行2020---2024年五年间的年度报告数据，运用信息不对称理论、全面风险管理理论及巴塞尔协议框架对其信用风险管理实践进行深入分析。研究发现：邯郸银行信用风险管理取得了一定成效，资产规模从2020年末1,809亿元增至2024年末2,492亿元，不良贷款率由2022年高点1.90%攀升至2023年2.24%后于2024年显著回落至1.38%，拨备覆盖率于2024年提升至204.35%，资本充足率持续高于监管要求；但同时也存在信贷结构高度集中于制造业与批发零售业、关注类贷款占比持续攀升（2023年达10.07%）、风险识别技术体系较为落后、风控执行机制存在松弛现象、信用风险人才相对匮乏等突出问题。

针对上述问题，本文提出优化信贷全流程管理、升级智能风控系统、强化执行层面管理与推进人才能力建设四个维度的对策建议，旨在为邯郸银行及同类城市商业银行提升信用风险管理水平提供参考。'''

para = doc.add_paragraph()
run = para.add_run(abstract_text)
run.font.size = Pt(12)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 关键词
kw_title = doc.add_paragraph()
run = kw_title.add_run('关键词：')
run.bold = True
run = kw_title.add_run('城市商业银行；信用风险；风险管理；邯郸银行；不良贷款率')

doc.add_paragraph()

# Abstract
abstract_title_en = doc.add_paragraph()
abstract_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = abstract_title_en.add_run('Abstract')
run.bold = True
run.font.size = Pt(14)

abstract_text_en = '''City commercial banks play a vital role in China's banking system by serving local economies and supporting small-to-medium enterprises. Due to their high regional concentration, heterogeneous borrower quality, and limited risk diversification capacity, credit risk remains their core challenge. This paper takes Handan Bank Co., Ltd. as a case study, systematically analyzes its annual report data from 2020 to 2024, and applies theoretical frameworks including information asymmetry, enterprise risk management, and the Basel Accord to evaluate its credit risk management practices. The research finds that while Handan Bank has achieved notable progress---total assets growing from CNY 180.9 billion in 2020 to CNY 249.2 billion in 2024, and non-performing loan (NPL) ratio declining sharply from 2.24% in 2023 to 1.38% in 2024---persistent problems remain, including a concentrated credit portfolio, rising special-mention loans (peaking at 10.07% of total loans in 2023), outdated risk identification technology, lapses in risk control execution, and insufficient risk management talent. The paper proposes recommendations across four dimensions: optimizing the full credit life cycle, upgrading intelligent risk control systems, strengthening execution management, and building talent capacity.

Keywords: city commercial bank; credit risk; risk management; Handan Bank; non-performing loan ratio'''

para = doc.add_paragraph()
run = para.add_run(abstract_text_en)
run.font.size = Pt(11)
para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ========== 目录 ==========
doc.add_page_break()
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_title.add_run('目  录')
run.bold = True
run.font.size = Pt(16)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

toc_items = [
    ('一、引言', '1'),
    ('（一）研究背景', '1'),
    ('（二）研究意义', '2'),
    ('（三）文献综述', '3'),
    ('（四）相关概念', '4'),
    ('二、理论基础', '5'),
    ('（一）信息不对称理论', '5'),
    ('（二）全面风险管理理论', '5'),
    ('（三）巴塞尔协议（资本监管框架）', '5'),
    ('三、邯郸银行信用风险管理的现状', '6'),
    ('（一）城市商业银行信用风险管理总体概况', '6'),
    ('（二）城市商业银行信用风险管理行业实践', '7'),
    ('（三）邯郸银行信用风险管理现状', '8'),
    ('（四）邯郸银行信用风险管理成效', '9'),
    ('【本节图表分析】', '10'),
    ('四、邯郸银行信用风险管理存在的问题', '11'),
    ('（一）信用风险管理体系与业务特点不匹配', '11'),
    ('（二）风险识别技术体系落后', '12'),
    ('（三）风险管控执行机制失效', '13'),
    ('（四）人才与内控建设滞后', '14'),
    ('【本节图表分析】', '14'),
    ('五、完善邯郸银行信用风险管理的对策建议', '15'),
    ('（一）优化信贷全流程', '15'),
    ('（二）升级智能风控系统', '16'),
    ('【大模型时代AI技术应用】', '17'),
    ('（三）强化执行层面管理', '18'),
    ('（四）人才能力建设工程', '19'),
    ('结论', '20'),
    ('参考文献', '21'),
]

for text, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'{text}\t\t{page}')
    run.font.size = Pt(12)

# 保存进度
doc.save('handan-bank-risk-paper-v2_part1.docx')
print("第一部分已保存")

# 继续创建完整的论文内容
# 这里我会使用一个更高效的方式，直接从markdown转换并添加图表